# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# 페이지 여백 설정 (2cm)
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# 폰트 기본 스타일
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif level == 2:
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return h

def add_screen_image(doc, filename, caption=""):
    img_path = os.path.join('e:/poject/OverseasWeb/docs_output/images', filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(10)
            rc = cp.add_run(f"▲ [화면 안내 스크린샷] {caption}")
            rc.font.name = '맑은 고딕'
            rc.font.size = Pt(9)
            rc.font.italic = True
            rc.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def add_callout(doc, text, title="💡 참고사항", bg_color="F0F7FF", border_color="3B82F6"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.8)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(title + "\n")
    run_t.font.name = '맑은 고딕'
    run_t.font.bold = True
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    
    run_c = p.add_run(text)
    run_c.font.name = '맑은 고딕'
    run_c.font.size = Pt(9.5)
    run_c.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# =============================================================
# 표지
# =============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(70)
title_p.paragraph_format.space_after = Pt(10)

run_sub = title_p.add_run("해외선교부 포털 시스템\n")
run_sub.font.name = '맑은 고딕'
run_sub.font.size = Pt(16)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

run_main = title_p.add_run("사용자 매뉴얼 & 실무 화면 가이드북")
run_main.font.name = '맑은 고딕'
run_main.font.size = Pt(26)
run_main.font.bold = True
run_main.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_p.paragraph_format.space_before = Pt(10)
desc_p.paragraph_format.space_after = Pt(110)
run_desc = desc_p.add_run("실제 시스템 UI 스크린샷과 단계별 조작 방법을 담은 사역자 표준 지침서")
run_desc.font.name = '맑은 고딕'
run_desc.font.size = Pt(11.5)
run_desc.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.paragraph_format.space_after = Pt(2)
run_info = info_p.add_run("• 발행 버전 : v2.0 (스크린샷 수록 완결본)\n• 발행 연도 : 2026년\n• 대상 : 전도담당자, 교역자, 선교부 사역자")
run_info.font.name = '맑은 고딕'
run_info.font.size = Pt(10)
run_info.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

doc.add_page_break()

# =============================================================
# 목차
# =============================================================
add_header_styled(doc, "목 차 (Table of Contents)", level=1)

toc_items = [
    ("1. 최초 로그인 방법", "로그인 화면 UI, 아이디/임시비밀번호 입력 및 OTP 팝업 조작"),
    ("2. 로그인 후 해야하는 것들", "프로필 화면 UI, 비밀번호 즉시 변경 및 담당 범위 점검"),
    ("3. OTP설정 방법", "회원정보 텔레그램 연동 폼, 봇 검색 및 Chat ID 자동 연동"),
    ("4. 전도담당자로써, 주간보고 작성방법", "주간보고 화면 UI, 5대 부서별 전도 지표 기입, 사진 첨부 및 최종 제출"),
    ("5. 계획 작성 방법", "전도 계획 탭 UI, 새 블록 추가, 세부 전략 작성 및 저장"),
    ("6. 월간보고 작성방법", "월간보고 탭 UI, 활동인원/사역자수 입력, 자동 연산, 수정 승인 요청"),
    ("7. 교회별 데이터 확인 방법", "종합 진단 대시보드 UI, 글로벌 KPI, Funnel 퍼널 분석, 추이 그래프"),
    ("부록. 자주 묻는 질문 (FAQ)", "로그인 실패, 텔레그램 미수신, 마감 후 수정 등 문제 해결")
]

for idx, (t, d) in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{t} ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    r2 = p.add_run(f"- {d}")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# =============================================================
# 1. 최초 로그인 방법
# =============================================================
add_header_styled(doc, "1. 최초 로그인 방법", level=1)
doc.add_paragraph("해외선교부 포털 시스템(Overseas Web Portal)에 접속하여 로그인하는 화면 및 조작 절차입니다.")

add_screen_image(doc, "screen_01_login.png", "로그인 페이지 및 2단계 OTP 인증 팝업 화면")

add_header_styled(doc, "1.1 화면 번호별 상세 조작 방법", level=2)
doc.add_paragraph("【화면 [1]번】 아이디 및 초기 임시 비밀번호 입력 :\n"
                  "• 관리자로부터 전달받은 본인의 사역자 아이디(예: missionary_tokyo)를 입력합니다.\n"
                  "• 함께 발급된 초기 임시 비밀번호를 비밀번호 칸에 정확히 입력합니다.\n\n"
                  "【화면 [2]번】 [로그인] 버튼 클릭 :\n"
                  "• 계정 정보 입력 후 파란색 [로그인 (Login)] 버튼을 클릭하거나 엔터키를 누릅니다.\n\n"
                  "【화면 [3]번】 텔레그램 6자리 OTP 인증번호 입력 :\n"
                  "• 2차 인증 보안 계정의 경우 화면 우측에 '2차 OTP 인증번호 입력' 팝업이 활성화됩니다.\n"
                  "• 텔레그램으로 전송된 6자리 번호를 입력하고 [인증 완료 및 로그인]을 클릭합니다. (제한시간 3분)")

add_callout(doc, "초기 비밀번호로 첫 로그인 시 '초기 계정 로그인에 성공하였습니다. 안전한 시스템 이용을 위해 먼저 비밀번호를 변경해주시기 바랍니다.' 팝업이 뜨며 비밀번호 변경 화면으로 자동 이동합니다.", title="⚠️ 초기 비밀번호 교체 안내", bg_color="FEF2F2", border_color="EF4444")

# =============================================================
# 2. 로그인 후 해야하는 것들
# =============================================================
add_header_styled(doc, "2. 로그인 후 해야하는 것들", level=1)
doc.add_paragraph("로그인 직후 보안을 강화하고 원활한 업무 환경을 구축하기 위해 필수적으로 수행해야 하는 체크리스트 화면입니다.")

add_screen_image(doc, "screen_02_profile_password.png", "회원 정보(프로필) 및 비밀번호 변경 화면")

add_header_styled(doc, "2.1 화면 번호별 상세 조작 방법", level=2)
doc.add_paragraph("【화면 [1]번】 소속 교회 및 배정 범위 점검 :\n"
                  "• 좌측 [기본 계정 정보] 카드에서 본인의 성명, 소속 권한 그룹(ROLE_USER / ROLE_ADMIN)을 확인합니다.\n"
                  "• [데이터 접근 범위]에 본인이 담당하는 국가 및 교회(예: 일본 / 도쿄교회, 오사카교회)가 올바르게 지정되어 있는지 확인합니다.\n\n"
                  "【화면 [2]번】 새 비밀번호 2회 입력 :\n"
                  "• 우측 [비밀번호 변경] 카드에서 [현재 비밀번호]를 입력한 후, 새로 사용할 [새 비밀번호]와 [새 비밀번호 확인]을 동일하게 입력합니다. (영문/숫자 4자리 이상)\n\n"
                  "【화면 [3]번】 [비밀번호 변경] 버튼 클릭 :\n"
                  "• 하단의 보라색 [비밀번호 변경 완료] 버튼을 누르면 새 비밀번호로 교체되며, 이후 로그인부터 새 암호가 적용됩니다.")

doc.add_page_break()

# =============================================================
# 3. OTP설정 방법
# =============================================================
add_header_styled(doc, "3. OTP설정 방법 (2단계 보안 인증)", level=1)
doc.add_paragraph("사역 데이터 보호를 위해 텔레그램 봇과 계정을 1회 연동하여 2차 인증(OTP) 번호 및 중요 공지를 수신하는 설정 화면입니다.")

add_screen_image(doc, "screen_03_otp_telegram.png", "텔레그램 연동 폼 및 봇 연동 테스트 화면")

add_header_styled(doc, "3.1 화면 번호별 상세 조작 방법", level=2)
doc.add_paragraph("【화면 [1]번】 Telegram ID (@사용자명) 입력 및 저장 :\n"
                  "• 좌측 [텔레그램 연동 설정] 카드에서 본인의 텔레그램 아이디(예: @missionary_tokyo)를 입력합니다.\n"
                  "• 하단의 [연동 정보 저장] 버튼을 클릭합니다.\n\n"
                  "【화면 [2]번】 텔레그램 앱에서 봇 대화방 시작 (/start) :\n"
                  "• 스마트폰 또는 PC의 텔레그램 앱에서 시스템 연동 봇(예: @OverseasPortalBot)을 검색합니다.\n"
                  "• 대화방 하단의 [시작] 또는 /start 를 전송하면, 봇이 아이디를 인식하여 [Telegram Chat ID]를 자동 연동합니다.\n\n"
                  "【화면 [3]번】 [테스트 메시지 발송]으로 수신 확인 :\n"
                  "• 우측 하단 [테스트 메시지 발송] 버튼을 클릭하여 본인의 텔레그램으로 테스트 알림이 도착하는지 최종 검증합니다.")

add_callout(doc, "자동 연동이 원활하지 않은 경우, 텔레그램 봇에게 /myid 메시지를 전송하여 답장받은 숫자 챗 ID를 포털의 [Telegram Chat ID] 입력란에 직접 넣고 저장할 수 있습니다.", title="💡 Chat ID 수동 입력 팁", bg_color="F0FDF4", border_color="10B981")

# =============================================================
# 4. 전도담당자로써, 주간보고 작성방법
# =============================================================
add_header_styled(doc, "4. 전도담당자로써, 주간보고 작성방법", level=1)
doc.add_paragraph("전도담당자가 매주 부서별 전도 현황을 입력하고, 현장 사진 및 특이사항을 작성하여 본부에 제출하는 핵심 화면입니다.")

add_screen_image(doc, "screen_04_weekly_report.png", "주간보고 작성 및 부서별 전도 지표 입력 화면")

add_header_styled(doc, "4.1 화면 번호별 상세 조작 방법", level=2)
doc.add_paragraph("【화면 [1]번】 교회 및 보고 주차 선택 :\n"
                  "• 상단 필터바에서 본인 [담당 교회]와 작성할 [보고 주차(예: 8월 3주차)]를 선택합니다.\n\n"
                  "【화면 [2]번】 5대 부서별 전도 지표 입력 :\n"
                  "• 교역자, 자문회, 장년회, 부녀회, 청년회의 지표를 표에 직접 입력합니다.\n"
                  "• 기입 항목: [전도재적], [찾기/탈락], [복음방/탈락], [가개강/탈락], [시온입교], [수료] 등 (합계 행 자동 계산)\n\n"
                  "【화면 [3]번】 사진 첨부 및 주간 특이사항 기입 :\n"
                  "• 우측 [현장 활동 사진] 영역을 클릭하여 전도 모임, 심방, 야외 전도 사진을 업로드합니다.\n"
                  "• [주간 특이사항 / 기도제목] 입력란에 사역 주요 성과 및 요청사항을 작성합니다.\n\n"
                  "【화면 [4]번】 임시 저장 및 최종 제출 :\n"
                  "• 작성 중간에는 [임시 저장] 버튼으로 데이터를 보관할 수 있습니다.\n"
                  "• 검토가 완료되면 초록색 [최종 제출] 버튼을 클릭하여 보고를 완료합니다. (매주 일요일 24:00 마감)")

doc.add_page_break()

# =============================================================
# 5. 계획 작성 방법
# =============================================================
add_header_styled(doc, "5. 계획 작성 방법 (전도 계획 수립)", level=1)
doc.add_paragraph("교회별 연간/분기별 전도 목표와 세부 실천 전략을 수립하고 체계적으로 관리하는 화면입니다.")

add_screen_image(doc, "screen_05_plan_tab.png", "전도관리 > 전도 계획 (Plan) 작성 화면")

add_header_styled(doc, "5.1 화면 번호별 상세 조작 방법", level=2)
doc.add_paragraph("【화면 [1]번】 [전도 계획 (Plan)] 탭 이동 :\n"
                  "• 상단 메뉴 [전도관리] 접속 후 3번째 서브탭인 [전도 계획 (Plan)]을 클릭합니다.\n\n"
                  "【화면 [2]번】 [+ 새 계획 블록 추가] 클릭 :\n"
                  "• 상단 우측의 [+ 새 계획 블록 추가] 파란색 버튼을 클릭하여 새로운 전략 카드 블록을 생성합니다.\n\n"
                  "【화면 [3]번】 전략 제목 및 세부 실행안 작성 :\n"
                  "• [계획 제목] : 핵심 전략 명칭 (예: 2026 하반기 대학가 집중 복음 전파 캠페인) 입력\n"
                  "• [세부 내용] : 목표 대상, 실행 방안, 투입 사역자, 기대 효과 등을 구체적으로 기입합니다.\n\n"
                  "【화면 [4]번】 [계획 저장 완료] 클릭 :\n"
                  "• 초록색 [💾 계획 저장 완료] 버튼을 누르면 서버에 즉시 반영되며, 최종 수정 일시와 작성자 아이디가 기록됩니다.")

# =============================================================
# 6. 월간보고 작성방법
# =============================================================
add_header_styled(doc, "6. 월간보고 작성방법", level=1)
doc.add_paragraph("한 달 동안의 전도 실적과 실활동 인원을 최종 결산하고, 전년 동월 대비 성장세를 분석하는 화면입니다.")

add_screen_image(doc, "screen_06_monthly_tab.png", "전도관리 > 월간보고 결산 작성 화면")

add_header_styled(doc, "6.1 화면 번호별 상세 조작 방법", level=2)
doc.add_paragraph("【화면 [1]번】 보고 연도 및 월 선택 :\n"
                  "• 상단 컨트롤 바에서 보고 대상 [연도]와 [월(1월~12월)]을 선택합니다.\n\n"
                  "【화면 [2]번】 실활동 인원수 및 전도 사역자수 입력 :\n"
                  "• [수정 모드]를 누른 뒤 5대 부서별 [실활동 인원수]와 [전도 사역자(인도자수)]를 기입합니다.\n\n"
                  "【화면 [3]번】 활동률(%) 및 전년 대비 증감율 자동 연산 확인 :\n"
                  "• 시스템이 [전도재적 대비 활동률] 및 [전년 동월 대비 증감율]을 실시간으로 자동 산출합니다.\n\n"
                  "【화면 [4]번】 과거 마감 데이터 [수정 권한 요청] :\n"
                  "• 이미 마감된 지난 달 데이터를 수정해야 하는 경우, [🔒 수정 권한 요청] 버튼을 눌러 관리자에게 승인 요청 사유를 전송합니다.")

doc.add_page_break()

# =============================================================
# 7. 교회별 데이터 확인 방법
# =============================================================
add_header_styled(doc, "7. 교회별 데이터 확인 방법 (종합 진단 대시보드)", level=1)
doc.add_paragraph("전 세계 교회들의 전도 추세, 부서별 기여도, 퍼널(Funnel) 단계별 전환율을 종합적으로 분석하는 화면입니다.")

add_screen_image(doc, "screen_07_diagnosis_dashboard.png", "홈 / 종합 진단 대시보드 화면")

add_header_styled(doc, "7.1 화면 번호별 상세 분석 방법", level=2)
doc.add_paragraph("【화면 [1]번】 글로벌 KPI 현황 요약 카드 :\n"
                  "• 상단 카드에서 전 세계 총 전도재적, 당주차 복음방 개설수, 시온입교 달성 실적을 한눈에 파악합니다.\n\n"
                  "【화면 [2]번】 단계별 전도 Funnel 퍼널 전환율 분석 :\n"
                  "• [찾기 → 복음방 → 가개강 → 시온입교] 전 과정의 전환율(%)과 단계별 이탈(Drop) 현황을 시각적으로 진단합니다.\n\n"
                  "【화면 [3]번】 다중 교회 주차별 비교 추이 그래프 (Trend) :\n"
                  "• 도쿄교회, 오사카교회, 나고야교회 등 여러 거점 교회의 주차별 성장 곡선을 상호 비교 분석합니다.")

# =============================================================
# 부록. 자주 묻는 질문 (FAQ)
# =============================================================
add_header_styled(doc, "부록. 자주 묻는 질문 (FAQ) 및 문제 해결", level=1)

faq_data = [
    ("Q1. 텔레그램으로 OTP 인증번호가 오지 않습니다.",
     "A1. ① [회원 정보]의 Telegram ID가 본인의 실제 @아이디와 일치하는지 확인하세요.\n"
     "② 텔레그램 앱에서 포털 전용 봇에게 /start 또는 /myid 메시지를 정상 전송했는지 확인하세요.\n"
     "③ 프로필 페이지에서 [테스트 발송]을 눌러 정상 수신 여부를 점검하세요."),
    ("Q2. 주간보고 제출 후 오타를 발견했는데 수정이 안 됩니다.",
     "A2. 주간보고는 당주차 일요일 자정(24:00)에 자동 마감됩니다. 마감 전에는 언제든 재수정/제출이 가능하지만, 마감 후에는 상급 관리자에게 문의하여 잠금 해제를 요청하셔야 합니다."),
    ("Q3. 담당하지 않는 타 교회의 데이터는 왜 조회가 안 되나요?",
     "A3. 시스템 보안 정책상 각 사역자는 본인에게 배정된 국가 및 담당 교회의 데이터만 접근할 수 있습니다. 추가 조회가 필요한 경우 총괄 관리자에게 권한 확장을 요청하세요.")
]

for q, a in faq_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    rq = p.add_run(q + "\n")
    rq.font.bold = True
    rq.font.size = Pt(10.5)
    rq.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    
    ra = p.add_run(a)
    ra.font.size = Pt(9.5)
    ra.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

out_docx = 'e:/poject/OverseasWeb/docs_output/해외선교부_포털_사용자_가이드.docx'
doc.save(out_docx)
print(f"Word document with UI screenshots successfully updated: {out_docx}")
